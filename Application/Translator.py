import os
import fitz  # PyMuPDF for PDF processing
import pandas as pd
from docx import Document
from docx.shared import Inches
import pytesseract
from PIL import Image
import io
import base64
from google import genai
from dotenv import load_dotenv
import re
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.lib.units import mm

load_dotenv()

# Configure Gemini API
client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))

# Add at the top of the file, after imports
SECONDARY_LANG = "Hindi"  # Change to 'Japanese' for final release

class FileTranslator:
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.xlsx', '.txt']
        
    def extract_text_from_pdf(self, pdf_path, use_ocr=False):
        """Extract text from PDF with optional OCR for images"""
        text_content = []
        
        try:
            # Use pdfplumber for better text extraction with structure preservation
            import pdfplumber
            
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    # Extract text with positioning information
                    text = page.extract_text()
                
                # If no text found and OCR is enabled, try OCR
                if not text.strip() and use_ocr:
                        # Convert page to image for OCR
                        img = page.to_image()
                        ocr_text = pytesseract.image_to_string(img.original, lang='jpn+eng')
                        text = ocr_text
                
                text_content.append(text)
            
            return "\n".join(text_content)
            
        except Exception as e:
            raise Exception(f"Error processing PDF: {str(e)}")
    
    def pdf_to_html_with_structure(self, pdf_path):
        """Convert PDF to HTML while preserving structure using pdfplumber"""
        try:
            import pdfplumber
            from bs4 import BeautifulSoup
            
            html_content = []
            html_content.append("""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    body { font-family: Arial, sans-serif; margin: 20px; }
                    .page { page-break-after: always; margin-bottom: 20px; }
                    .text-block { margin: 5px 0; }
                    .paragraph { margin: 10px 0; }
                </style>
            </head>
            <body>
            """)
            
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    html_content.append(f'<div class="page" id="page-{page_num + 1}">')
                    
                    # Extract text blocks with positioning
                    text_blocks = page.extract_text_blocks()
                    
                    if text_blocks:
                        for block in text_blocks:
                            if block.get('text', '').strip():
                                # Preserve some formatting based on font size
                                font_size = block.get('size', 12)
                                if font_size > 14:
                                    html_content.append(f'<h2 class="text-block">{block["text"]}</h2>')
                                elif font_size > 12:
                                    html_content.append(f'<h3 class="text-block">{block["text"]}</h3>')
                                else:
                                    html_content.append(f'<p class="text-block">{block["text"]}</p>')
                    else:
                        # Fallback to simple text extraction
                        text = page.extract_text()
                        if text.strip():
                            paragraphs = text.split('\n\n')
                            for para in paragraphs:
                                if para.strip():
                                    html_content.append(f'<p class="paragraph">{para.strip()}</p>')
                    
                    html_content.append('</div>')
            
            html_content.append("</body></html>")
            return '\n'.join(html_content)
            
        except Exception as e:
            raise Exception(f"Error converting PDF to HTML: {str(e)}")
    
    def translate_html_content(self, html_content, source_lang, target_lang):
        """Translate text content within HTML while preserving structure"""
        try:
            from bs4 import BeautifulSoup
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Translate text nodes but preserve HTML structure
            for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'span', 'div']):
                if element.string and element.string.strip():
                    try:
                        translated_text = self.translate_text(element.string.strip(), source_lang, target_lang)
                        element.string = translated_text
                    except Exception as e:
                        # Keep original text if translation fails
                        continue
            
            return str(soup)
            
        except Exception as e:
            raise Exception(f"Error translating HTML content: {str(e)}")
    
    def html_to_pdf_with_fonts(self, html_content, output_path, target_lang=None):
        """Convert HTML to PDF with proper font support"""
        try:
            from weasyprint import HTML, CSS
            from weasyprint.text.fonts import FontConfiguration
            
            # Configure fonts based on target language
            if target_lang and (target_lang.lower() == 'hindi' or target_lang.lower() == 'hin'):
                # Use Devanagari font for Hindi
                css_content = """
                @font-face {
                    font-family: 'Noto Sans Devanagari';
                    src: url('Tiro_Devanagari_Hindi/NotoSansDevanagari-Regular.ttf') format('truetype');
                }
                body { 
                    font-family: 'Noto Sans Devanagari', Arial, sans-serif; 
                    margin: 20px; 
                }
                .page { page-break-after: always; margin-bottom: 20px; }
                .text-block { margin: 5px 0; }
                .paragraph { margin: 10px 0; }
                """
            else:
                # Use Japanese font for Japanese
                css_content = """
                @font-face {
                    font-family: 'Noto Sans JP';
                    src: url('Noto_Sans_JP/NotoSansJP-VariableFont_wght.ttf') format('truetype');
                }
                body { 
                    font-family: 'Noto Sans JP', Arial, sans-serif; 
                    margin: 20px; 
                }
                .page { page-break-after: always; margin-bottom: 20px; }
                .text-block { margin: 5px 0; }
                .paragraph { margin: 10px 0; }
                """
            
            # Configure font
            font_config = FontConfiguration()
            css = CSS(string=css_content, font_config=font_config)
            
            # Convert HTML to PDF
            HTML(string=html_content).write_pdf(
                output_path,
                stylesheets=[css],
                font_config=font_config
            )
            
        except Exception as e:
            raise Exception(f"Error converting HTML to PDF: {str(e)}")
    
    def extract_text_from_docx(self, docx_path):
        """Extract text from Word document"""
        try:
            doc = Document(docx_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)
            
            return "\n".join(text_content)
            
        except Exception as e:
            raise Exception(f"Error processing Word document: {str(e)}")
    
    def extract_text_from_excel(self, excel_path):
        """Extract text from Excel file"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(excel_path)
            text_content = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_path, sheet_name=sheet_name)
                text_content.append(f"Sheet: {sheet_name}")
                text_content.append(df.to_string(index=False))
                text_content.append("\n")
            
            return "\n".join(text_content)
            
        except Exception as e:
            raise Exception(f"Error processing Excel file: {str(e)}")
    
    def extract_text_from_file(self, file_path, use_ocr=False):
        """Extract text from any supported file format"""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path, use_ocr)
        elif file_extension == '.docx':
            return self.extract_text_from_docx(file_path)
        elif file_extension == '.xlsx':
            return self.extract_text_from_excel(file_path)
        elif file_extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            raise Exception(f"Unsupported file format: {file_extension}")
    
    def translate_text(self, text, source_lang, target_lang):
        """Translate text using Gemini API"""
        try:
            # Determine language codes
            lang_map = {
                'Japanese': 'Japanese',
                'Hindi': 'Hindi',
                'English': 'English',
                'JAP': 'Japanese',
                'HIN': 'Hindi',
                'ENG': 'English'
            }
            
            source = lang_map.get(source_lang, source_lang)
            target = lang_map.get(target_lang, target_lang)
            
            prompt = f"""
            Translate the following text from {source} to {target}. 
            Maintain the original formatting, structure, and meaning as much as possible.
            If the text contains technical terms, preserve them appropriately.
            
            Text to translate:
            {text}
            
            Provide only the translated text without any additional explanations.
            """
            
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=prompt
            )
            
            return response.text.strip()
            
        except Exception as e:
            raise Exception(f"Translation error: {str(e)}")

    def save_translated_pdf(self, original_path, translated_text, output_path, font_path=None, preserve_structure=True, source_lang=None, target_lang=None):
        """Save translated text as PDF, preserving structure if requested. Use Hindi or Japanese font as needed."""
        if preserve_structure and source_lang and target_lang:
            try:
                # Convert PDF to HTML with structure
                html_content = self.pdf_to_html_with_structure(original_path)
                
                # Translate the HTML content
                translated_html = self.translate_html_content(html_content, source_lang, target_lang)
                
                # Convert back to PDF with proper fonts
                self.html_to_pdf_with_fonts(translated_html, output_path, target_lang)
                
            except Exception as e:
                # Fallback to simple text approach
                self.save_translated_pdf_simple(translated_text, output_path, target_lang)
        else:
            self.save_translated_pdf_simple(translated_text, output_path, target_lang)
    
    def save_translated_pdf_simple(self, translated_text, output_path, target_lang=None):
        """Simple PDF generation without structure preservation"""
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.units import mm
            
            # Select font based on target_lang
            if target_lang and (target_lang.lower() == 'hindi' or target_lang.lower() == 'hin'):
                font_path = "Tiro_Devanagari_Hindi/NotoSansDevanagari-Regular.ttf"
                font_name = "NotoSansDevanagariRegular"
            else:
                font_path = "Noto_Sans_JP/NotoSansJP-VariableFont_wght.ttf"
                font_name = "NotoSansJP"
            
                font_path = os.path.abspath(font_path)
                if not os.path.exists(font_path):
                    print(f"DEBUG: Font file not found at {font_path}")
                    raise Exception(f"Font file not found: {font_path}.")
            
                # Register font only if not already registered
                if font_name not in pdfmetrics.getRegisteredFontNames():
                    pdfmetrics.registerFont(TTFont(font_name, font_path))
            
                c = canvas.Canvas(output_path, pagesize=A4)
                width, height = A4
                c.setFont(font_name, 12)
                margin_x = 20 * mm
                margin_y = 20 * mm
                line_height = 15
                x = margin_x
                y = height - margin_y
            
                for line in translated_text.split('\n'):
                    if y - line_height < margin_y:
                        c.showPage()
                        c.setFont(font_name, 12)
                        y = height - margin_y
                    c.drawString(x, y, line)
                    y -= line_height
            
                c.save()
            
        except Exception as e:
            raise Exception(f"Error saving PDF: {str(e)}")
    
    def save_translated_docx(self, original_path, translated_text, output_path, target_lang=None):
        """Save translated text as Word document, using Devanagari font for Hindi output."""
        try:
            doc = Document()
            # Split text into paragraphs
            paragraphs = translated_text.split('\n\n')
            # Choose font for Hindi
            use_hindi_font = target_lang and (target_lang.lower() == 'hindi' or target_lang.lower() == 'hin')
            hindi_font_name = "Mangal"  # Or use 'Noto Sans Devanagari' if installed
            for para in paragraphs:
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    if use_hindi_font:
                        for run in p.runs:
                            run.font.name = hindi_font_name
            doc.save(output_path)
        except Exception as e:
            raise Exception(f"Error saving Word document: {str(e)}")
    
    def save_translated_excel(self, original_path, translated_text, output_path):
        """Save translated text as Excel file"""
        try:
            # Simple approach: save as single sheet with translated text
            df = pd.DataFrame({'Translated_Content': [translated_text]})
            df.to_excel(output_path, index=False)
            
        except Exception as e:
            raise Exception(f"Error saving Excel file: {str(e)}")
    
    def save_translated_file(self, original_path, translated_text, output_path, **kwargs):
        """Save translated text in the same format as original"""
        file_extension = os.path.splitext(original_path)[1].lower()
        if file_extension == '.pdf':
            self.save_translated_pdf(original_path, translated_text, output_path, **kwargs)
        elif file_extension == '.docx':
            # Pass target_lang if available
            target_lang = kwargs.get('target_lang', None)
            self.save_translated_docx(original_path, translated_text, output_path, target_lang=target_lang)
        elif file_extension == '.xlsx':
            self.save_translated_excel(original_path, translated_text, output_path)
        elif file_extension == '.txt':
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(translated_text)
        else:
            raise Exception(f"Unsupported file format for saving: {file_extension}")
    
    def translate_file(self, file_path, source_lang, target_lang, use_ocr=False, output_path=None):
        """Complete file translation process"""
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            print(f"Extracting text from {file_path}...")
            original_text = self.extract_text_from_file(file_path, use_ocr)
            if not original_text.strip():
                raise Exception("No text content found in the file")
            print(f"Translating from {source_lang} to {target_lang}...")
            translated_text = self.translate_text(original_text, source_lang, target_lang)
            if output_path is None:
                base_name = os.path.splitext(file_path)[0]
                extension = os.path.splitext(file_path)[1]
                output_path = f"{base_name}_translated_{target_lang}{extension}"
            if file_extension == '.pdf':
                self.save_translated_file(file_path, translated_text, output_path, preserve_structure=True, source_lang=source_lang, target_lang=target_lang)
            else:
                self.save_translated_file(file_path, translated_text, output_path)
            return {
                'original_text': original_text,
                'translated_text': translated_text,
                'output_path': output_path,
                'success': True
            }
        except Exception as e:
            return {
                'error': str(e),
                'success': False
            }

# Utility functions
def detect_language(text):
    """Detect the language of the text"""
    try:
        prompt = f"""
        Detect the language of the following text. Respond with only the language name (e.g., "Japanese", "English").
        
        Text: {text[:500]}  # Use first 500 characters for detection
        
        Language:
        """
        
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt
        )
        
        return response.text.strip()
        
    except Exception as e:
        return "Unknown"

def validate_file_format(file_path):
    """Validate if file format is supported"""
    supported_formats = ['.pdf', '.docx', '.xlsx', '.txt']
    file_extension = os.path.splitext(file_path)[1].lower()
    return file_extension in supported_formats
